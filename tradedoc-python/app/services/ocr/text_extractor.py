"""
文档文字提取器 — 免费、精准、优先于 OCR

支持：
  - PDF（电子版）→ pdfplumber 直接提取文字和表格
  - Word (.docx) → python-docx 提取
  - Excel (.xlsx) → 后续扩展

只有提取失败或内容为空（扫描件）时，才回退到 OCR
"""

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


class TextExtractor:
    """文档文字提取器"""

    def extract(self, file_path: str) -> dict:
        """
        提取文件文字内容

        Returns:
            {
                "success": True/False,
                "text": "提取的全部文字",
                "tables": [{"headers": [], "rows": [[]]}],  # 表格数据
                "pages": 1,
                "method": "pdfplumber" / "python-docx" / "none",
            }
        """
        ext = Path(file_path).suffix.lower()

        if ext == ".pdf":
            return self._extract_pdf(file_path)
        elif ext in (".docx", ".doc"):
            return self._extract_docx(file_path)
        elif ext in (".jpg", ".jpeg", ".png", ".tiff", ".tif"):
            # 图片无法直接提取文字，需要 OCR
            return {"success": False, "text": "", "tables": [], "pages": 0, "method": "none"}
        else:
            return {"success": False, "text": "", "tables": [], "pages": 0, "method": "none"}

    def _extract_pdf(self, file_path: str) -> dict:
        """PDF 文字提取（pdfplumber）"""
        try:
            import pdfplumber

            all_text = []
            all_tables = []

            with pdfplumber.open(file_path) as pdf:
                pages = len(pdf.pages)

                for i, page in enumerate(pdf.pages[:10]):  # 最多10页
                    # 提取文字
                    text = page.extract_text()
                    if text:
                        all_text.append(text)

                    # 提取表格
                    tables = page.extract_tables()
                    for table in tables:
                        if table and len(table) > 1:
                            headers = [str(c or "") for c in table[0]]
                            rows = [[str(c or "") for c in row] for row in table[1:]]
                            all_tables.append({
                                "page": i + 1,
                                "headers": headers,
                                "rows": rows,
                            })

            full_text = "\n".join(all_text)

            # 判断是否有效提取（扫描件提取的文字会很少或为空）
            if len(full_text.strip()) < 20:
                logger.info(f"PDF 文字提取内容过少（{len(full_text)}字），可能是扫描件")
                return {"success": False, "text": full_text, "tables": all_tables, "pages": pages, "method": "pdfplumber"}

            logger.info(f"PDF 文字提取成功: {len(full_text)}字, {len(all_tables)}个表格, {pages}页")
            return {
                "success": True,
                "text": full_text,
                "tables": all_tables,
                "pages": pages,
                "method": "pdfplumber",
            }

        except ImportError:
            logger.error("pdfplumber 未安装")
            return {"success": False, "text": "", "tables": [], "pages": 0, "method": "none"}
        except Exception as e:
            logger.error(f"PDF 提取失败: {e}")
            return {"success": False, "text": "", "tables": [], "pages": 0, "method": "none"}

    def _extract_docx(self, file_path: str) -> dict:
        """Word 文档文字提取（python-docx）"""
        try:
            from docx import Document

            doc = Document(file_path)

            # 提取段落
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            full_text = "\n".join(paragraphs)

            # 提取表格
            tables = []
            for table in doc.tables:
                headers = [cell.text for cell in table.rows[0].cells] if table.rows else []
                rows = [[cell.text for cell in row.cells] for row in table.rows[1:]]
                tables.append({"headers": headers, "rows": rows})

            if len(full_text.strip()) < 10:
                return {"success": False, "text": full_text, "tables": tables, "pages": 1, "method": "python-docx"}

            logger.info(f"Word 提取成功: {len(full_text)}字, {len(tables)}个表格")
            return {
                "success": True,
                "text": full_text,
                "tables": tables,
                "pages": 1,
                "method": "python-docx",
            }

        except ImportError:
            logger.error("python-docx 未安装")
            return {"success": False, "text": "", "tables": [], "pages": 0, "method": "none"}
        except Exception as e:
            logger.error(f"Word 提取失败: {e}")
            return {"success": False, "text": "", "tables": [], "pages": 0, "method": "none"}


# 全局单例
text_extractor = TextExtractor()
